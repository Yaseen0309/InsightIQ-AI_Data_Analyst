import os
import re
from docx import Document
import pandas as pd
import pdfplumber
from PIL import Image
import pytesseract
from modules.data_utils import normalize_columns

def load_file(file_path):
    ext = os.path.splitext(file_path)[1].lower().strip('.')

    # PDF text extraction
    if ext == 'pdf':
        with pdfplumber.open(file_path) as pdf:
            return '\n'.join([page.extract_text() or '' for page in pdf.pages]), None

    # Plain text file
    elif ext == 'txt':
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read(), None
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='ISO-8859-1') as f:
                return f.read(), None

    # Word Document (.docx)
    elif ext in ['doc', 'docx']:
        doc_result, _ = extract_text_from_docx(file_path)
        if isinstance(doc_result, pd.DataFrame):
            doc_result, col_map = normalize_columns(doc_result)
            return doc_result, col_map
        else:
            return doc_result, None

    elif ext in ['png', 'jpg', 'jpeg', 'tiff', 'bmp']:
        img_result, _ = extract_text_from_image(file_path)
        if isinstance(img_result, pd.DataFrame):
            img_result, col_map = normalize_columns(img_result)
            return img_result, col_map
        else:
            return img_result, None

    # CSV or Excel
    elif ext in ['csv', 'xlsx']:
        try:
            if ext == 'csv':
                try:
                    df = pd.read_csv(file_path)
                except UnicodeDecodeError:
                    print("⚠️ UTF-8 failed, trying ISO-8859-1...")
                    df = pd.read_csv(file_path, encoding='ISO-8859-1')
            else:
                df = pd.read_excel(file_path)

            df, col_map = normalize_columns(df)
            return df, col_map

        except Exception as e:
            print("❌ Error reading tabular file:", e)
            raise e

    else:
        print(f"🚨 DEBUG: Unsupported extension .{ext}")
        raise ValueError(f"Unsupported file type: .{ext}")




def extract_text_from_image(file_path):
    try:
        image = Image.open(file_path).convert("RGB")
        raw_text = pytesseract.image_to_string(image)

        lines = [line.strip() for line in raw_text.split("\n") if line.strip()]
        if len(lines) < 2:
            return raw_text, None

        # Step 1: Attempt to split lines into columns (tab or 2+ spaces)
        rows = []
        for line in lines:
            if "\t" in line:
                rows.append(line.split("\t"))
            else:
                # Split on 2+ spaces (common in Excel screenshots)
                split_line = re.split(r"\s{2,}", line)
                if len(split_line) > 1:
                    rows.append(split_line)

        if len(rows) < 2:
            return raw_text, None  # Not enough structure for a table

        # Step 2: Filter rows with consistent length
        col_counts = [len(r) for r in rows]
        most_common_len = max(set(col_counts), key=col_counts.count)
        table_rows = [r for r in rows if len(r) == most_common_len]

        if len(table_rows) >= 2:
            df = pd.DataFrame(table_rows[1:], columns=table_rows[0])
            return df, None

        return raw_text, None  # Fall back if reconstruction fails

    except Exception as e:
        return f"❌ Failed to extract text from image: {e}", None

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        tables = doc.tables

        if tables:
            # Support multiple tables in the future (for now, first one)
            for table in tables:
                data = []
                for row in table.rows:
                    data.append([cell.text.strip() for cell in row.cells])
                if len(data) >= 2:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    return df, None
            # If no usable table found
            return "⚠️ Tables found but not parseable as structured data.", None
        else:
            full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            return '\n'.join(full_text), None

    except Exception as e:
        return f"❌ Failed to read .docx file: {e}", None



